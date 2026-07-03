"""
Extracteur DOCX — Extraction de texte depuis des fichiers Microsoft Word.
Utilise python-docx.
"""

from pathlib import Path


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extrait le texte complet d'un fichier DOCX (Word).

    Args:
        file_path: Chemin vers le fichier .docx.

    Returns:
        Texte extrait du document.

    Raises:
        FileNotFoundError: Si le fichier n'existe pas.
        ValueError: Si le fichier n'est pas un .docx valide.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {file_path}")

    if file_path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Le fichier n'est pas un document Word : {file_path}")

    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx n'est pas installé. Exécutez : pip install python-docx"
        )

    doc = Document(str(file_path))

    text_parts: list[str] = []

    # Extraire les paragraphes
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extraire les tableaux
    for table_idx, table in enumerate(doc.tables, start=1):
        table_text = _format_docx_table(table, table_idx)
        if table_text:
            text_parts.append(table_text)

    return "\n\n".join(text_parts)


def _format_docx_table(table, table_idx: int) -> str:
    """
    Formate un tableau DOCX en texte lisible.

    Args:
        table: Objet Table de python-docx.
        table_idx: Index du tableau dans le document.

    Returns:
        Représentation textuelle du tableau.
    """
    lines: list[str] = [f"--- Tableau {table_idx} ---"]

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))

    return "\n".join(lines) if len(lines) > 1 else ""
